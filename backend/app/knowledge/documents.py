from __future__ import annotations

import csv
import json
import re
import tempfile
import zipfile
from dataclasses import dataclass, field
from html import unescape
from pathlib import Path
from typing import Any, Dict, List, Union
from xml.etree import ElementTree


@dataclass
class ParsedDocument:
    path: Path
    title: str
    text: str
    metadata: Dict[str, Any] = field(default_factory=dict)


class DocumentParser:
    """Parse documents into normalized plain text."""

    DIRECT_TEXT_EXTENSIONS = {
        ".txt",
        ".md",
        ".rst",
        ".py",
        ".js",
        ".ts",
        ".java",
        ".cpp",
        ".c",
        ".h",
        ".css",
        ".scss",
        ".toml",
        ".ini",
        ".cfg",
        ".conf",
        ".yaml",
        ".yml",
        ".log",
        ".sql",
    }

    STRUCTURED_EXTENSIONS = {".json", ".csv", ".html", ".htm", ".docx", ".pdf", ".doc", ".ppt", ".pptx", ".xls", ".xlsx"}

    def discover_files(self, inputs: List[str], recursive: bool = True) -> List[Path]:
        files: List[Path] = []
        for raw in inputs:
            path = Path(raw).expanduser()
            if not path.exists():
                continue
            if path.is_file():
                files.append(path.resolve())
                continue
            iterator = path.rglob("*") if recursive else path.glob("*")
            for candidate in iterator:
                if candidate.is_file() and self.is_supported(candidate):
                    files.append(candidate.resolve())
        return sorted(set(files))

    def is_supported(self, path: Path) -> bool:
        return path.suffix.lower() in (self.DIRECT_TEXT_EXTENSIONS | self.STRUCTURED_EXTENSIONS)

    def parse(self, path: Union[str, Path]) -> ParsedDocument:
        file_path = Path(path).expanduser().resolve()
        suffix = file_path.suffix.lower()
        text = ""

        if suffix in self.DIRECT_TEXT_EXTENSIONS:
            text = self._read_text_file(file_path)
        elif suffix == ".json":
            text = self._read_json(file_path)
        elif suffix == ".csv":
            text = self._read_csv(file_path)
        elif suffix in {".html", ".htm"}:
            text = self._read_html(file_path)
        elif suffix == ".docx":
            text = self._read_docx(file_path)
        elif suffix == ".pdf":
            text = self._read_pdf(file_path)
        elif suffix in {".doc", ".ppt", ".pptx", ".xls", ".xlsx"}:
            text = self._read_with_markitdown(file_path)
        else:
            text = self._read_text_file(file_path)

        cleaned = self._normalize_text(text)
        return ParsedDocument(
            path=file_path,
            title=file_path.name,
            text=cleaned,
            metadata={"suffix": suffix, "size": file_path.stat().st_size},
        )

    def _read_text_file(self, path: Path) -> str:
        for encoding in ("utf-8", "utf-8-sig", "gbk", "latin-1"):
            try:
                return path.read_text(encoding=encoding)
            except UnicodeDecodeError:
                continue
        return path.read_text(encoding="utf-8", errors="ignore")

    def _read_json(self, path: Path) -> str:
        try:
            payload = json.loads(self._read_text_file(path))
            return json.dumps(payload, ensure_ascii=False, indent=2)
        except json.JSONDecodeError:
            return self._read_text_file(path)

    def _read_csv(self, path: Path) -> str:
        rows: List[str] = []
        with path.open("r", encoding="utf-8", errors="ignore", newline="") as handle:
            reader = csv.reader(handle)
            for row in reader:
                rows.append(" | ".join(cell.strip() for cell in row))
        return "\n".join(rows)

    def _read_html(self, path: Path) -> str:
        content = self._read_text_file(path)
        try:
            from bs4 import BeautifulSoup

            soup = BeautifulSoup(content, "html.parser")
            return soup.get_text("\n")
        except Exception:
            without_script = re.sub(r"<script[\s\S]*?</script>", "", content, flags=re.IGNORECASE)
            without_style = re.sub(r"<style[\s\S]*?</style>", "", without_script, flags=re.IGNORECASE)
            no_tags = re.sub(r"<[^>]+>", " ", without_style)
            return unescape(no_tags)

    def _read_docx(self, path: Path) -> str:
        try:
            from docx import Document

            document = Document(str(path))
            paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
            return "\n".join(paragraphs)
        except Exception:
            pass

        with zipfile.ZipFile(path) as archive:
            xml_content = archive.read("word/document.xml")
        root = ElementTree.fromstring(xml_content)
        paragraphs = []
        namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        for paragraph in root.findall(".//w:p", namespace):
            texts = [node.text for node in paragraph.findall(".//w:t", namespace) if node.text]
            line = "".join(texts).strip()
            if line:
                paragraphs.append(line)
        return "\n".join(paragraphs)

    def _read_pdf(self, path: Path) -> str:
        text = self._read_pdf_with_pypdf(path)
        if text:
            return text

        text = self._read_pdf_with_ppstructure(path)
        if text:
            return text

        text = self._read_with_markitdown(path)
        if text:
            return text

        raise RuntimeError(
            f"Unable to parse PDF {path.name}. "
            "Install pypdf for text PDFs, or install paddleocr and pymupdf for scanned PDFs."
        )

    def _read_pdf_with_pypdf(self, path: Path) -> str:
        try:
            from pypdf import PdfReader

            reader = PdfReader(str(path))
            pages = []
            for index, page in enumerate(reader.pages, start=1):
                page_text = page.extract_text() or ""
                if page_text.strip():
                    pages.append(f"# Page {index}\n{page_text.strip()}")
            return "\n\n".join(pages)
        except Exception:
            return ""

    def _read_pdf_with_ppstructure(self, path: Path) -> str:
        try:
            from paddleocr import PPStructure
        except Exception:
            return ""

        table_engine = PPStructure(show_log=False, image_orientation=True)
        sections: List[str] = []
        with tempfile.TemporaryDirectory(prefix="pdf-ocr-") as temp_dir, tempfile.TemporaryDirectory(
            prefix="pdf-pages-"
        ) as pages_dir:
            output_dir = Path(temp_dir)
            page_images = self._render_pdf_pages(path, Path(pages_dir))
            if not page_images:
                return ""

            try:
                from paddleocr.ppstructure.utility import save_structure_res
            except Exception:
                save_structure_res = None

            for index, image_path in enumerate(page_images, start=1):
                result = table_engine(str(image_path)) or []
                page_text = self._ppstructure_result_to_markdown(result).strip()
                if not page_text:
                    continue
                if save_structure_res is not None:
                    save_structure_res(result, str(output_dir), f"{path.stem}_page_{index}")
                sections.append(f"# Page {index}\n{page_text}")
        return "\n\n".join(sections)

    def _render_pdf_pages(self, path: Path, output_dir: Path) -> List[Path]:
        try:
            import fitz
        except Exception:
            return []

        image_paths: List[Path] = []
        document = fitz.open(str(path))
        try:
            for index, page in enumerate(document, start=1):
                pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
                image_path = output_dir / f"{path.stem}_page_{index}.png"
                pixmap.save(str(image_path))
                image_paths.append(image_path)
        finally:
            document.close()
        return image_paths

    def _ppstructure_result_to_markdown(self, result: List[Dict[str, Any]]) -> str:
        blocks: List[str] = []
        for block in result:
            block_type = str(block.get("type", "")).lower()
            if "html" in block and block["html"]:
                blocks.append(str(block["html"]).strip())
                continue

            if block_type == "table":
                text = str(block.get("res", "")).strip()
                if text:
                    blocks.append(text)
                continue

            raw_res = block.get("res")
            if isinstance(raw_res, list):
                lines = []
                for item in raw_res:
                    if isinstance(item, dict):
                        candidate = item.get("text")
                        if candidate and str(candidate).strip():
                            lines.append(str(candidate).strip())
                if lines:
                    blocks.append("\n".join(lines))
                continue

            if isinstance(raw_res, dict):
                candidate = raw_res.get("text")
                if candidate and str(candidate).strip():
                    blocks.append(str(candidate).strip())
        return "\n\n".join(block for block in blocks if block.strip())

    def _read_with_markitdown(self, path: Path) -> str:
        try:
            from markitdown import MarkItDown

            result = MarkItDown().convert(str(path))
            text = getattr(result, "text_content", "") or ""
            return str(text)
        except Exception:
            return ""

    def _normalize_text(self, text: str) -> str:
        normalized = (text or "").replace("\r\n", "\n").replace("\r", "\n")
        normalized = re.sub(r"\n{3,}", "\n\n", normalized)
        normalized = "\n".join(line.rstrip() for line in normalized.splitlines())
        return normalized.strip()
